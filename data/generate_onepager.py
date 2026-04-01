"""
TFAIS Database — One-Pager Generator
Writes data/tfais_db_onepager.docx
"""
import os, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

OUTPUT = Path(__file__).resolve().parent / "tfais_db_onepager.docx"

# (column, type, what it means)
TABLES = [
    ("districts", "Stores the 37 Tamil Nadu districts.", [
        ("code",     "TEXT",  "District ID used by the portal (e.g. '20')"),
        ("name_ta",  "TEXT",  "District name in Tamil"),
    ]),
    ("blocks", "Administrative blocks inside each district.", [
        ("code",        "TEXT", "Block ID from the portal"),
        ("name_ta",     "TEXT", "Block name in Tamil"),
        ("district_id", "FK",  "→ districts.id"),
    ]),
    ("dealers", "Fertilizer dealers, one row per shop.", [
        ("dealer_code", "TEXT", "Licence/registration number (empty if not on card)"),
        ("name_ta",     "TEXT", "Dealer name in Tamil"),
        ("contact",     "TEXT", "Phone number"),
        ("block_id",    "FK",  "→ blocks.id"),
    ]),
    ("fertilizer_stock", "Core fact table — daily stock per dealer per fertilizer.", [
        ("dealer_id",       "FK",    "→ dealers.id"),
        ("fertilizer_name", "TEXT",  "Fertilizer type (Tamil name from card header)"),
        ("quantity",        "FLOAT", "Available stock"),
        ("unit",            "TEXT",  "Unit of measure (default: KG)"),
        ("scrape_date",     "DATE",  "Date the data represents"),
        ("scrape_run_id",   "FK",    "→ scrape_runs.id"),
    ]),
    ("scrape_runs", "One row per pipeline execution — audit log.", [
        ("status",          "TEXT", "running / completed / failed / partial"),
        ("trigger_type",    "TEXT", "manual / scheduled / resume"),
        ("dealers_scraped", "INT",  "Dealer cards saved in this run"),
        ("errors_count",    "INT",  "Recoverable errors hit during the run"),
        ("started_at",      "TS",   "When the run began"),
        ("completed_at",    "TS",   "When it finished (NULL if still running)"),
    ]),
    ("scrape_checkpoints", "Per-(district, block) progress tracker for resume-on-failure.", [
        ("scrape_run_id",  "FK",   "→ scrape_runs.id"),
        ("district_code",  "TEXT", "District being scraped"),
        ("block_code",     "TEXT", "Block being scraped"),
        ("status",         "TEXT", "pending / done / error"),
        ("dealers_found",  "INT",  "Cards found for this block"),
    ]),
]

HEADER_COLOR = "1F497D"   # dark blue
ALT_ROW      = "DCE6F1"   # light blue
WHITE        = "FFFFFF"

def shade(cell, hex_color):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pr.append(shd)

def set_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Inches(widths[i])

def add_mini_table(doc, rows):
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"

    # header
    hdr = t.rows[0].cells
    for cell, txt in zip(hdr, ("Column", "Type", "Description")):
        cell.text = txt
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, HEADER_COLOR)

    for i, (col, typ, desc) in enumerate(rows):
        r = t.add_row().cells
        r[0].text, r[1].text, r[2].text = col, typ, desc
        for cell in r:
            cell.paragraphs[0].runs[0].font.size = Pt(8)
        if i % 2 == 0:
            for cell in r:
                shade(cell, ALT_ROW)

    set_widths(t, [1.5, 0.55, 3.45])
    return t


def build():
    doc = Document()

    # ── page margins (narrow so table fits) ──────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
        setattr(section, attr, Cm(1.5))

    # ── title ─────────────────────────────────────────────────────────────────
    h = doc.add_heading("TFAIS — Database Quick Reference", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    h.runs[0].font.size = Pt(16)

    sub = doc.add_paragraph("Tamil Nadu Fertilizer Availability Intelligence System  •  PostgreSQL  •  6 Tables")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
    sub.paragraph_format.space_after = Pt(10)

    # ── one block per table ───────────────────────────────────────────────────
    for name, desc, cols in TABLES:
        # table name as a compact heading
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # one-line description
        d = doc.add_paragraph(desc)
        d.paragraph_format.space_before = Pt(0)
        d.paragraph_format.space_after  = Pt(3)
        d.runs[0].font.size = Pt(8.5)
        d.runs[0].font.italic = True
        d.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        add_mini_table(doc, cols)

    # ── footer note ───────────────────────────────────────────────────────────
    from datetime import date
    note = doc.add_paragraph(f"\nOnly key columns shown.  Auto-generated {date.today()}.")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].font.size = Pt(7.5)
    note.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
