"""
TFAIS Database Documentation & Export Script
=============================================
Generates:
  1. data/tfais_db_documentation.docx  — Word doc with table-by-table schema docs
  2. data/tfais_db_export.xlsx         — Excel file with one sheet per table (live data)

Requirements:
    pip install python-docx openpyxl sqlalchemy psycopg2-binary python-dotenv pandas
"""

import os
import sys
from pathlib import Path

# Allow running from any directory
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from dotenv import load_dotenv
load_dotenv(PROJECT_ROOT / ".env")

import pandas as pd
from sqlalchemy import create_engine, inspect, text
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# DB connection
# ---------------------------------------------------------------------------
DB_URL = os.getenv("DATABASE_URL") or (
    f"postgresql://{os.getenv('DB_USER', 'postgres')}:{os.getenv('DB_PASSWORD', 'postgres')}"
    f"@{os.getenv('DB_HOST', 'localhost')}:{os.getenv('DB_PORT', '5432')}/{os.getenv('DB_NAME', 'tfais')}"
)

engine = create_engine(DB_URL)

# ---------------------------------------------------------------------------
# Schema metadata — authoritative descriptions derived from models.py
# ---------------------------------------------------------------------------
TABLE_META = {
    "districts": {
        "description": (
            "Reference/dimension table storing all Tamil Nadu districts. "
            "Each row represents one district whose code is used as a foreign key "
            "by the blocks table. Populated once during the first scrape run."
        ),
        "columns": {
            "id":         ("INTEGER", "Auto-incrementing surrogate primary key."),
            "code":       ("VARCHAR(20)", "Unique district code as used by the government portal (e.g., '20'). Indexed."),
            "name_ta":    ("VARCHAR(200)", "District name in Tamil, as scraped from the portal."),
            "created_at": ("TIMESTAMPTZ", "UTC timestamp when this row was first inserted."),
        },
    },
    "blocks": {
        "description": (
            "Reference/dimension table for blocks (administrative subdivisions) within each district. "
            "A block belongs to exactly one district. The composite unique constraint "
            "(code, district_id) prevents duplicate blocks across runs."
        ),
        "columns": {
            "id":          ("INTEGER", "Auto-incrementing surrogate primary key."),
            "code":        ("VARCHAR(20)", "Block code as used by the portal. Indexed."),
            "name_ta":     ("VARCHAR(200)", "Block name in Tamil, as scraped."),
            "district_id": ("INTEGER FK → districts.id", "Foreign key linking this block to its parent district."),
            "created_at":  ("TIMESTAMPTZ", "UTC timestamp when this row was first inserted."),
        },
    },
    "dealers": {
        "description": (
            "Reference/dimension table for individual fertilizer dealers. "
            "Each dealer belongs to one block. A partial unique index on "
            "(dealer_code, block_id) enforces deduplication only when dealer_code is non-empty, "
            "accommodating cards that have no extractable code."
        ),
        "columns": {
            "id":          ("INTEGER", "Auto-incrementing surrogate primary key."),
            "dealer_code": ("VARCHAR(50)", "Dealer licence/registration code (may be empty string if not present on card). Indexed."),
            "name_ta":     ("VARCHAR(300)", "Dealer name in Tamil."),
            "address":     ("TEXT", "Dealer address as scraped (nullable)."),
            "contact":     ("VARCHAR(20)", "Contact phone number (nullable)."),
            "block_id":    ("INTEGER FK → blocks.id", "Foreign key linking this dealer to their block."),
            "created_at":  ("TIMESTAMPTZ", "UTC timestamp of first insert."),
            "updated_at":  ("TIMESTAMPTZ", "UTC timestamp of last update (auto-refreshed on upsert)."),
        },
    },
    "fertilizer_stock": {
        "description": (
            "Core fact/time-series table. Each row records the available stock quantity "
            "of one fertilizer type at one dealer on one scrape date. "
            "The unique constraint (dealer_id, fertilizer_name, scrape_date) prevents "
            "duplicate rows if the scraper re-runs on the same day. "
            "fertilizer_name is stored as-is from the Tamil card headers "
            "(no separate master table) to avoid mapping complexity."
        ),
        "columns": {
            "id":              ("BIGINT", "Auto-incrementing surrogate primary key (BigInt for long-term scalability)."),
            "dealer_id":       ("INTEGER FK → dealers.id", "Foreign key to the dealer whose stock this row records."),
            "fertilizer_name": ("VARCHAR(100)", "Fertilizer type name in Tamil, taken directly from card column header."),
            "quantity":        ("FLOAT", "Stock quantity available (in units defined by the 'unit' column)."),
            "unit":            ("VARCHAR(10)", "Unit of measurement (default: 'KG')."),
            "scrape_date":     ("DATE", "Logical date this stock snapshot represents (separate from created_at)."),
            "created_at":      ("TIMESTAMPTZ", "UTC timestamp when this row was written to the database."),
            "scrape_run_id":   ("INTEGER FK → scrape_runs.id", "Foreign key to the scrape run that produced this row (nullable for legacy data)."),
        },
    },
    "scrape_runs": {
        "description": (
            "Audit/tracking table with one row per end-to-end pipeline execution. "
            "Captures run status, trigger type, and aggregate counts for observability "
            "and resume-on-failure logic."
        ),
        "columns": {
            "id":               ("INTEGER", "Auto-incrementing primary key; used as run identifier across all related tables."),
            "started_at":       ("TIMESTAMPTZ", "UTC timestamp when this run began."),
            "completed_at":     ("TIMESTAMPTZ", "UTC timestamp when the run ended (NULL while still running)."),
            "status":           ("VARCHAR(20)", "Run lifecycle state: 'running' | 'completed' | 'failed' | 'partial'."),
            "trigger_type":     ("VARCHAR(20)", "What initiated this run: 'manual' | 'scheduled' | 'resume'."),
            "districts_total":  ("INTEGER", "Total number of districts targeted in this run (nullable)."),
            "blocks_total":     ("INTEGER", "Total number of blocks targeted in this run (nullable)."),
            "dealers_scraped":  ("INTEGER", "Number of dealer cards successfully parsed and saved."),
            "errors_count":     ("INTEGER", "Number of recoverable errors encountered during the run."),
            "notes":            ("TEXT", "Free-text notes or error summary (nullable)."),
        },
    },
    "scrape_checkpoints": {
        "description": (
            "Granular checkpoint log recording the scrape status of every "
            "(district, block) pair within a run. Enables resume-on-failure: "
            "a new run skips pairs already marked 'done'. "
            "Uses string codes instead of FK integer IDs because checkpoints are written "
            "before the district/block rows are guaranteed to be committed."
        ),
        "columns": {
            "id":             ("INTEGER", "Auto-incrementing primary key."),
            "scrape_run_id":  ("INTEGER FK → scrape_runs.id", "Foreign key to the parent run."),
            "district_code":  ("VARCHAR(20)", "District code string (matches districts.code)."),
            "block_code":     ("VARCHAR(20)", "Block code string (matches blocks.code within the district)."),
            "status":         ("VARCHAR(20)", "Checkpoint state: 'pending' | 'done' | 'error'."),
            "dealers_found":  ("INTEGER", "Number of dealer cards found for this (district, block) pair."),
            "error_message":  ("TEXT", "Error details if status='error' (nullable)."),
            "completed_at":   ("TIMESTAMPTZ", "UTC timestamp when this checkpoint was last updated (nullable)."),
        },
    },
}

TABLE_ORDER = [
    "districts",
    "blocks",
    "dealers",
    "fertilizer_stock",
    "scrape_runs",
    "scrape_checkpoints",
]

OUTPUT_DIR = Path(__file__).resolve().parent
DOCX_PATH  = OUTPUT_DIR / "tfais_db_documentation.docx"
XLSX_PATH  = OUTPUT_DIR / "tfais_db_export.xlsx"


# ---------------------------------------------------------------------------
# Helper: shade a table row
# ---------------------------------------------------------------------------
def _shade_row(row, hex_color: str):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def _set_col_widths(table, widths_inches):
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_inches):
                cell.width = Inches(widths_inches[idx])


# ---------------------------------------------------------------------------
# Generate Word documentation
# ---------------------------------------------------------------------------
def generate_word_doc():
    doc = Document()

    # ---- Title page ----
    title = doc.add_heading("TFAIS Database Documentation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph("Tamil Nadu Fertilizer Availability Intelligence System")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)

    doc.add_paragraph(f"Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')} IST")
    doc.add_page_break()

    # ---- Overview ----
    doc.add_heading("Database Overview", level=1)
    doc.add_paragraph(
        "The TFAIS PostgreSQL database is structured around a simple star-like schema. "
        "Three dimension tables (districts, blocks, dealers) capture the geographic hierarchy. "
        "One fact table (fertilizer_stock) stores time-series stock snapshots. "
        "Two operational tables (scrape_runs, scrape_checkpoints) track pipeline execution "
        "and enable resume-on-failure."
    )

    # ---- Table count summary ----
    doc.add_heading("Tables at a Glance", level=2)
    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = "Light List Accent 1"
    hdr = summary_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Table", "Type", "Primary Purpose"
    _shade_row(summary_table.rows[0], "1F497D")
    for cell in summary_table.rows[0].cells:
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True

    summary_rows = [
        ("districts",           "Dimension",   "Tamil Nadu districts reference list"),
        ("blocks",              "Dimension",   "Blocks within each district"),
        ("dealers",             "Dimension",   "Fertilizer dealer master data"),
        ("fertilizer_stock",    "Fact",        "Daily stock snapshots per dealer/fertilizer"),
        ("scrape_runs",         "Operational", "Pipeline run audit log"),
        ("scrape_checkpoints",  "Operational", "Per-(district, block) resume checkpoints"),
    ]
    for i, (tbl, typ, purpose) in enumerate(summary_rows):
        row = summary_table.add_row().cells
        row[0].text, row[1].text, row[2].text = tbl, typ, purpose
        if i % 2 == 0:
            _shade_row(summary_table.rows[i + 1], "DCE6F1")

    doc.add_paragraph()

    # ---- Per-table sections ----
    for tbl_name in TABLE_ORDER:
        meta = TABLE_META[tbl_name]

        doc.add_heading(f"Table: {tbl_name}", level=1)
        doc.add_paragraph(meta["description"])

        col_table = doc.add_table(rows=1, cols=3)
        col_table.style = "Light List Accent 1"
        hdr = col_table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Column", "Data Type", "Description"
        _shade_row(col_table.rows[0], "1F497D")
        for cell in col_table.rows[0].cells:
            run = cell.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True

        for i, (col_name, (dtype, desc)) in enumerate(meta["columns"].items()):
            row = col_table.add_row().cells
            row[0].text = col_name
            row[1].text = dtype
            row[2].text = desc
            if i % 2 == 0:
                _shade_row(col_table.rows[i + 1], "DCE6F1")

        _set_col_widths(col_table, [1.5, 2.0, 3.0])
        doc.add_paragraph()

    doc.save(DOCX_PATH)
    print(f"[OK] Word doc saved: {DOCX_PATH}")


# ---------------------------------------------------------------------------
# Generate Excel export
# ---------------------------------------------------------------------------
def generate_excel_export():
    with engine.connect() as conn:
        with pd.ExcelWriter(XLSX_PATH, engine="openpyxl") as writer:
            for tbl_name in TABLE_ORDER:
                try:
                    df = pd.read_sql_table(tbl_name, conn)
                    # Excel cannot handle timezone-aware datetimes — strip tz info
                    for col in df.select_dtypes(include=["datetimetz"]).columns:
                        df[col] = df[col].dt.tz_localize(None)
                    df.to_excel(writer, sheet_name=tbl_name, index=False)

                    # Auto-fit column widths
                    ws = writer.sheets[tbl_name]
                    for col_cells in ws.columns:
                        max_len = max(
                            (len(str(cell.value)) if cell.value is not None else 0)
                            for cell in col_cells
                        )
                        ws.column_dimensions[col_cells[0].column_letter].width = min(max_len + 4, 60)

                    print(f"  [OK] Sheet '{tbl_name}': {len(df):,} rows")
                except Exception as e:
                    print(f"  [WARN] Could not export '{tbl_name}': {e}")

    print(f"[OK] Excel export saved: {XLSX_PATH}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    print("=== TFAIS DB Documentation & Export ===\n")

    print("1/2 Generating Word documentation...")
    generate_word_doc()

    print("\n2/2 Exporting database to Excel...")
    generate_excel_export()

    print("\nDone. Files written to:", OUTPUT_DIR)
